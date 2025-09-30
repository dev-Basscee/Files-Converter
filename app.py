from flask import Flask, request, render_template, send_file, jsonify
import os
from werkzeug.utils import secure_filename
import io
from PIL import Image
import PyPDF2
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import img2pdf
from pdf2image import convert_from_path
import zipfile
import json

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

# Supported conversions mapping
CONVERSIONS = {
    # Image formats
    'png': ['jpg', 'jpeg', 'pdf', 'bmp', 'gif', 'tiff', 'webp', 'ico'],
    'jpg': ['png', 'pdf', 'bmp', 'gif', 'tiff', 'webp', 'ico'],
    'jpeg': ['png', 'pdf', 'bmp', 'gif', 'tiff', 'webp', 'ico'],
    'bmp': ['png', 'jpg', 'jpeg', 'pdf', 'gif', 'tiff', 'webp'],
    'gif': ['png', 'jpg', 'jpeg', 'pdf', 'bmp', 'tiff', 'webp'],
    'tiff': ['png', 'jpg', 'jpeg', 'pdf', 'bmp', 'gif', 'webp'],
    'webp': ['png', 'jpg', 'jpeg', 'pdf', 'bmp', 'gif', 'tiff'],
    'ico': ['png', 'jpg', 'jpeg', 'pdf', 'bmp'],
    
    # Document formats
    'pdf': ['txt', 'png', 'jpg', 'jpeg'],
    'txt': ['pdf', 'docx'],
    'docx': ['txt', 'pdf'],
    
    # Add more format support as needed
}

def get_file_extension(filename):
    """Get file extension without the dot."""
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

def convert_image(input_path, output_format):
    """Convert image from one format to another."""
    img = Image.open(input_path)
    
    # Handle transparency for formats that don't support it
    if output_format.lower() in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
        background = Image.new('RGB', img.size, (255, 255, 255))
        if img.mode == 'P':
            img = img.convert('RGBA')
        background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
        img = background
    
    output = io.BytesIO()
    
    # Special handling for different formats
    if output_format.lower() in ['jpg', 'jpeg']:
        img.convert('RGB').save(output, format='JPEG', quality=95)
    elif output_format.lower() == 'ico':
        img.save(output, format='ICO', sizes=[(256, 256)])
    else:
        img.save(output, format=output_format.upper())
    
    output.seek(0)
    return output

def convert_image_to_pdf(input_path):
    """Convert image to PDF."""
    output = io.BytesIO()
    with open(input_path, 'rb') as f:
        output.write(img2pdf.convert(f.read()))
    output.seek(0)
    return output

def convert_pdf_to_text(input_path):
    """Extract text from PDF."""
    output = io.BytesIO()
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        output.write('\n\n'.join(text).encode('utf-8'))
    output.seek(0)
    return output

def convert_pdf_to_image(input_path, output_format='png'):
    """Convert PDF first page to image."""
    images = convert_from_path(input_path, first_page=1, last_page=1)
    output = io.BytesIO()
    if images:
        if output_format.lower() in ['jpg', 'jpeg']:
            images[0].convert('RGB').save(output, format='JPEG', quality=95)
        else:
            images[0].save(output, format=output_format.upper())
    output.seek(0)
    return output

def convert_text_to_pdf(input_path):
    """Convert text file to PDF."""
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as file:
        text = file.read()
    
    width, height = letter
    y = height - 50
    
    for line in text.split('\n'):
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line[:100])  # Limit line length
        y -= 15
    
    c.save()
    output.seek(0)
    return output

def convert_text_to_docx(input_path):
    """Convert text file to DOCX."""
    output = io.BytesIO()
    doc = Document()
    
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as file:
        for line in file:
            doc.add_paragraph(line.rstrip())
    
    doc.save(output)
    output.seek(0)
    return output

def convert_docx_to_text(input_path):
    """Convert DOCX to text."""
    output = io.BytesIO()
    doc = Document(input_path)
    
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    output.write('\n'.join(text).encode('utf-8'))
    output.seek(0)
    return output

def convert_docx_to_pdf(input_path):
    """Convert DOCX to PDF (simplified - just extract text and create PDF)."""
    output = io.BytesIO()
    doc = Document(input_path)
    
    c = canvas.Canvas(output, pagesize=letter)
    width, height = letter
    y = height - 50
    
    for paragraph in doc.paragraphs:
        if y < 50:
            c.showPage()
            y = height - 50
        text = paragraph.text[:100]  # Limit line length
        c.drawString(50, y, text)
        y -= 15
    
    c.save()
    output.seek(0)
    return output

def perform_conversion(input_path, input_format, output_format):
    """Perform the actual file conversion."""
    
    # Image to image conversions
    if input_format in ['png', 'jpg', 'jpeg', 'bmp', 'gif', 'tiff', 'webp', 'ico']:
        if output_format in ['png', 'jpg', 'jpeg', 'bmp', 'gif', 'tiff', 'webp', 'ico']:
            return convert_image(input_path, output_format)
        elif output_format == 'pdf':
            return convert_image_to_pdf(input_path)
    
    # PDF conversions
    elif input_format == 'pdf':
        if output_format == 'txt':
            return convert_pdf_to_text(input_path)
        elif output_format in ['png', 'jpg', 'jpeg']:
            return convert_pdf_to_image(input_path, output_format)
    
    # Text conversions
    elif input_format == 'txt':
        if output_format == 'pdf':
            return convert_text_to_pdf(input_path)
        elif output_format == 'docx':
            return convert_text_to_docx(input_path)
    
    # DOCX conversions
    elif input_format == 'docx':
        if output_format == 'txt':
            return convert_docx_to_text(input_path)
        elif output_format == 'pdf':
            return convert_docx_to_pdf(input_path)
    
    raise ValueError(f"Conversion from {input_format} to {output_format} is not supported")

@app.route('/')
def index():
    """Render the main page."""
    return render_template('index.html', conversions=CONVERSIONS)

@app.route('/api/supported-formats', methods=['GET'])
def supported_formats():
    """Get all supported conversion formats."""
    return jsonify(CONVERSIONS)

@app.route('/api/convert', methods=['POST'])
def convert_file():
    """Handle file conversion."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        output_format = request.form.get('output_format', '').lower()
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not output_format:
            return jsonify({'error': 'No output format specified'}), 400
        
        # Get input format
        input_format = get_file_extension(file.filename)
        
        if not input_format:
            return jsonify({'error': 'Invalid file format'}), 400
        
        # Check if conversion is supported
        if input_format not in CONVERSIONS:
            return jsonify({'error': f'Input format "{input_format}" is not supported'}), 400
        
        if output_format not in CONVERSIONS.get(input_format, []):
            return jsonify({'error': f'Cannot convert from "{input_format}" to "{output_format}"'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        
        # Perform conversion
        converted_data = perform_conversion(input_path, input_format, output_format)
        
        # Clean up uploaded file
        os.remove(input_path)
        
        # Generate output filename
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}.{output_format}"
        
        # Return converted file
        return send_file(
            converted_data,
            as_attachment=True,
            download_name=output_filename,
            mimetype=f'application/{output_format}'
        )
        
    except Exception as e:
        # Clean up files if they exist
        if 'input_path' in locals() and os.path.exists(input_path):
            os.remove(input_path)
        
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
